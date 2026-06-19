"""
utils/attachments.py  --" Attachment classification, saving, and LLM content-block assembly.

This is a pure utility module (not an agent module).  It has no tools, hooks, or
prompts.  Bridges and the gateway call it directly.

Pipeline
--------
1. classify(attachment) â†' sets Attachment.kind based on mime_type / filename extension
2. save_upload(attachment, uploads_dir) â†' writes bytes to disk, returns Path
3. build_content_blocks(text, attachments, model_cfg, att_cfg, workspace)
     â†' list[dict]  (OpenAI-compat content block list for a user message)

Content block strategies
------------------------
image/*      + vision model  â†' {"type": "image_url", "image_url": {"url": "data:<mime>;base64,..."}}
image/*      + no vision     â†' reference note in text, saved to uploads/
text/* + md/py/json/etc.     â†' {"type": "text", "text": "<fenced code block>"}
application/pdf              â†' {"type": "text", "text": "<extracted or stub>"}  (pdfplumber if available)
.docx                        â†' {"type": "text", "text": "<extracted>"}  (python-docx if available)
binary / unknown             â†' reference note in text, saved to uploads/

The inline-vs-reference decision respects AttachmentConfig thresholds:
  inline_max_files  --" max number of files to inline per message (default 3)
  inline_max_bytes  --" max total raw bytes to inline (default ~200 KB)
Once either threshold is hit, remaining files are reference-only regardless of kind.

Soft dependencies
-----------------
pdfplumber           --" PDF text extraction.  If absent, PDFs get a stub reference note.
python-docx          --" DOCX text extraction. If absent, DOCX files get a stub reference note.
rapidocr_onnxruntime --" OCR fallback for scanned PDF pages and embedded DOCX images.
                         If absent, those pages/images just fall back to whatever the
                         text layer already produced (no hard failure).
None are required; attachments.py degrades gracefully.
"""

from __future__ import annotations

import base64
import hashlib
import json
import logging
from pathlib import Path

from TinyCTX.contracts import Attachment, AttachmentKind
from TinyCTX.config import ModelConfig, AttachmentConfig

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Extension â†' MIME overrides for types mimetypes may not know
# ---------------------------------------------------------------------------
_EXT_MIME: dict[str, str] = {
    ".md":   "text/markdown",
    ".py":   "text/x-python",
    ".ts":   "text/typescript",
    ".tsx":  "text/typescript",
    ".jsx":  "text/javascript",
    ".yaml": "text/yaml",
    ".yml":  "text/yaml",
    ".toml": "text/toml",
    ".sh":   "text/x-shellscript",
    ".json": "application/json",
    ".csv":  "text/csv",
    ".xml":  "text/xml",
    ".html": "text/html",
    ".htm":  "text/html",
    ".sql":  "text/x-sql",
    ".rs":   "text/x-rust",
    ".go":   "text/x-go",
    ".java": "text/x-java",
    ".c":    "text/x-c",
    ".cpp":  "text/x-c++",
    ".h":    "text/x-c",
}

# Extensions we treat as text regardless of reported MIME
_TEXT_EXTENSIONS: frozenset[str] = frozenset(_EXT_MIME.keys()) | {
    ".txt", ".log", ".ini", ".cfg", ".conf", ".env",
    ".r", ".rb", ".php", ".swift", ".kt", ".cs",
}

# Image MIME prefixes we can inline
_IMAGE_MIMES: frozenset[str] = frozenset({
    "image/jpeg", "image/png", "image/gif", "image/webp",
})


# ---------------------------------------------------------------------------
# Classification
# ---------------------------------------------------------------------------

def classify(attachment: Attachment) -> AttachmentKind:
    """
    Derive the AttachmentKind for an Attachment from its mime_type and filename.
    The returned kind is not stored back onto the (frozen) dataclass  --" callers
    use the return value directly.
    """
    mime  = (attachment.mime_type or "").lower().split(";")[0].strip()
    ext   = Path(attachment.filename).suffix.lower()

    # Prefer extension overrides for text types (bridges often lie about MIME)
    if ext in _TEXT_EXTENSIONS:
        return AttachmentKind.TEXT

    if mime in _IMAGE_MIMES:
        return AttachmentKind.IMAGE

    if mime == "image/svg+xml":
        return AttachmentKind.TEXT  # SVG is XML text

    if mime == "application/pdf" or ext == ".pdf":
        return AttachmentKind.DOCUMENT

    if mime in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ) or ext == ".docx":
        return AttachmentKind.DOCUMENT

    if mime.startswith("text/"):
        return AttachmentKind.TEXT

    if mime == "application/json":
        return AttachmentKind.TEXT

    return AttachmentKind.BINARY


# ---------------------------------------------------------------------------
# Saving
# ---------------------------------------------------------------------------

def _load_cache(uploads_dir: Path) -> dict[str, str]:
    """Load uploads/cache.json â†' {sha256hex: filename}.  Returns {} on any error."""
    cache_path = uploads_dir / "cache.json"
    try:
        return json.loads(cache_path.read_text(encoding="utf-8"))
    except FileNotFoundError:
        return {}
    except Exception as exc:
        logger.warning("Could not read upload cache: %s", exc)
        return {}


def _save_cache(uploads_dir: Path, cache: dict[str, str]) -> None:
    """Persist cache.json atomically (write to .tmp, rename)."""
    cache_path = uploads_dir / "cache.json"
    tmp_path   = uploads_dir / "cache.json.tmp"
    try:
        tmp_path.write_text(json.dumps(cache, indent=2), encoding="utf-8")
        tmp_path.replace(cache_path)
    except Exception as exc:
        logger.warning("Could not write upload cache: %s", exc)


def save_upload(attachment: Attachment, uploads_dir: Path) -> Path:
    """
    Write attachment bytes to uploads_dir/<filename>, deduplicating by content.

    A sidecar uploads/cache.json maps sha256 â†' stored filename so that
    re-uploading the same file returns the existing path without a write.
    Original filenames are preserved; a counter suffix is only added when
    two genuinely different files share the same name.

    Creates the directory if it doesn't exist.
    Returns the absolute path of the saved file.
    If writing fails, logs the error and returns the intended path anyway
    so the caller can still include a reference note.
    """
    uploads_dir.mkdir(parents=True, exist_ok=True)
    # Sanitize filename: strip all directory components to prevent path traversal.
    safe_name = Path(attachment.filename).name or "unnamed"
    # Protect the dedup sidecar from being clobbered by an uploaded file.
    if safe_name in {"cache.json", "cache.json.tmp"}:
        safe_name = safe_name + ".upload"
    # Belt-and-suspenders: verify the resolved path stays inside uploads_dir.
    def _check(p: Path) -> Path:
        if not str(p.resolve()).startswith(str(uploads_dir.resolve())):
            logger.error("Attachment filename escapes uploads dir: %s", attachment.filename)
            raise ValueError(f"Filename escapes uploads directory: {attachment.filename}")
        return p

    content_hash = hashlib.sha256(attachment.data).hexdigest()
    cache = _load_cache(uploads_dir)

    # Cache hit  --" same bytes already stored.
    if content_hash in cache:
        existing = uploads_dir / cache[content_hash]
        if existing.exists():
            logger.debug("Dedup: reusing existing upload %s", existing.name)
            return existing
        # Stale entry (file was deleted)  --" fall through and re-save.
        logger.debug("Upload cache stale for %s, re-saving", safe_name)

    # Determine destination, avoiding name collisions with different-content files.
    dest = _check(uploads_dir / safe_name)
    if dest.exists():
        stem    = dest.stem
        suffix  = dest.suffix
        counter = 1
        while dest.exists():
            dest = _check(uploads_dir / f"{stem}_{counter}{suffix}")
            counter += 1

    try:
        dest.write_bytes(attachment.data)
        logger.debug("Saved upload: %s (%d bytes)", dest.name, len(attachment.data))
    except OSError as exc:
        logger.error("Failed to save upload %s: %s", dest, exc)
        return dest

    cache[content_hash] = dest.name
    _save_cache(uploads_dir, cache)
    return dest


# ---------------------------------------------------------------------------
# Image conversion helpers
# ---------------------------------------------------------------------------

def convert_to_png(data: bytes) -> bytes | None:
    """Convert image bytes to PNG using Pillow.  Returns None if unavailable.

    ICC profiles are explicitly stripped (icc_profile=None) because some model
    backends (llama.cpp / llama-swap) reject PNG input that contains an iCCP
    chunk, returning HTTP 400 "Failed to load image".
    """
    try:
        import io
        from PIL import Image
        img = Image.open(io.BytesIO(data))
        buf = io.BytesIO()
        img.save(buf, format="PNG", icc_profile=None)
        return buf.getvalue()
    except ImportError:
        return None
    except Exception as exc:
        logger.warning("Image conversion to PNG failed: %s", exc)
        return None


# ---------------------------------------------------------------------------
# Text extraction helpers (soft deps)
# ---------------------------------------------------------------------------

_MIN_CHARS_PER_PAGE = 20  # below this, treat a PDF page as scanned and try OCR
_ocr_engine = None        # lazily-constructed RapidOCR singleton (model load is expensive)


def _get_ocr_engine():
    """Lazily construct and cache the RapidOCR engine.  Raises ImportError if absent."""
    global _ocr_engine
    if _ocr_engine is None:
        from rapidocr_onnxruntime import RapidOCR
        _ocr_engine = RapidOCR()
    return _ocr_engine


def _ocr_image(img) -> str | None:
    """Run RapidOCR on a PIL image.  Returns None if rapidocr_onnxruntime is unavailable."""
    try:
        import numpy as np
        engine = _get_ocr_engine()
        result, _ = engine(np.array(img))
        return "\n".join(line[1] for line in result) if result else None
    except ImportError:
        return None
    except Exception as exc:
        logger.warning("RapidOCR failed: %s", exc)
        return None


def extract_pdf_text(data: bytes) -> str | None:
    """
    Extract text from a PDF using pdfplumber's text layer.  Pages with little or
    no extractable text (scanned/image pages) are rasterized and run through
    RapidOCR as a fallback.  Returns None if pdfplumber is unavailable.
    """
    try:
        import io
        import pdfplumber
    except ImportError:
        return None

    pages: list[str] = []
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                text = (page.extract_text() or "").strip()
                if len(text) < _MIN_CHARS_PER_PAGE:
                    ocr_text = _ocr_image(page.to_image(resolution=150).original)
                    if ocr_text:
                        text = ocr_text
                pages.append(text)
    except Exception as exc:
        logger.warning("PDF text extraction failed: %s", exc)
        return None
    return "\n\n".join(p for p in pages if p.strip()) or None


def extract_docx_text(data: bytes) -> str | None:
    """
    Extract text from a DOCX using python-docx: paragraphs, table cells, and
    OCR of embedded images (covers pasted screenshots / scanned pages).
    Returns None if python-docx is unavailable.
    """
    try:
        import io
        from docx import Document
    except ImportError:
        return None

    try:
        doc = Document(io.BytesIO(data))
        parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]

        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))

        for shape in doc.inline_shapes:
            try:
                from PIL import Image
                rId = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
                blob = doc.part.related_parts[rId].blob
                ocr_text = _ocr_image(Image.open(io.BytesIO(blob)))
                if ocr_text:
                    parts.append(ocr_text)
            except ImportError:
                break  # Pillow unavailable -- skip remaining images, keep text/tables
            except Exception:
                continue

        return "\n\n".join(parts) or None
    except Exception as exc:
        logger.warning("DOCX text extraction failed: %s", exc)
        return None


# ---------------------------------------------------------------------------
# Content block assembly
# ---------------------------------------------------------------------------

def build_content_blocks(
    text:        str,
    attachments: tuple[Attachment, ...],
    model_cfg:   ModelConfig,
    att_cfg:     AttachmentConfig,
    workspace:   Path,
) -> list[dict] | str:
    """
    Build an OpenAI-compat content block list for a user message that includes
    one or more attachments.

    Returns a list[dict] if any attachment is inlineable (so the caller must
    send content as a list rather than a plain string), or a plain str with
    reference notes appended if every attachment is reference-only.

    Args:
        text:        The user's text message.
        attachments: Tuple of Attachment objects from InboundMessage.
        model_cfg:   ModelConfig of the currently active LLM.
        att_cfg:     AttachmentConfig thresholds from Config.
        workspace:   Resolved workspace root Path.

    The function:
      1. Saves every attachment to workspace/uploads/ (permanent record).
      2. Decides inline vs reference per file, respecting thresholds.
      3. Builds content blocks for inlineable files.
      4. Appends reference notes for non-inlineable files.
    """
    if not attachments:
        return text

    uploads_dir = workspace / att_cfg.uploads_dir

    inline_count = 0
    inline_bytes = 0
    blocks: list[dict] = []
    ref_notes: list[str] = []

    for att in attachments:
        kind = classify(att)
        saved_path = save_upload(att, uploads_dir)

        # --- threshold check ---
        over_count = inline_count >= att_cfg.inline_max_files
        over_bytes = (inline_bytes + len(att.data)) > att_cfg.inline_max_bytes
        force_ref  = over_count or over_bytes

        # --- strategy selection ---
        if force_ref or kind == AttachmentKind.BINARY:
            ref_notes.append(
                f"[File uploaded to {saved_path}: {att.filename}]"
            )
            continue

        if kind == AttachmentKind.IMAGE:
            if not model_cfg.supports_vision:
                ref_notes.append(
                    f"[Image uploaded to {saved_path}: {att.filename} -- model does not support vision, use filesystem tools to inspect]"
                )
                continue
            # Inline as image_url block -- always convert to PNG for broad API compatibility
            img_data = att.data
            mime = att.mime_type.split(";")[0].strip()
            if mime != "image/png":
                converted = convert_to_png(img_data)
                if converted is not None:
                    img_data = converted
                    mime = "image/png"
                else:
                    ref_notes.append(
                        f"[Image uploaded to {saved_path}: {att.filename}"
                        " -- unsupported image format and Pillow unavailable for conversion]"
                    )
                    continue
            b64 = base64.b64encode(img_data).decode()
            blocks.append({
                "type": "image_url",
                "image_url": {"url": f"data:{mime};base64,{b64}"},
            })
            inline_count += 1
            inline_bytes += len(att.data)
            continue

        if kind == AttachmentKind.TEXT:
            try:
                content = att.data.decode("utf-8", errors="replace")
            except Exception:
                content = att.data.decode("latin-1", errors="replace")
            ext  = Path(att.filename).suffix.lstrip(".")
            lang = ext if ext else "text"
            block_text = f"**{att.filename}**\n```{lang}\n{content}\n```"
            blocks.append({"type": "text", "text": block_text})
            inline_count += 1
            inline_bytes += len(att.data)
            continue

        if kind == AttachmentKind.DOCUMENT:
            ext = Path(att.filename).suffix.lower()
            extracted: str | None = None

            if ext == ".pdf" or att.mime_type == "application/pdf":
                extracted = extract_pdf_text(att.data)
                if extracted is None:
                    ref_notes.append(
                        f"[PDF uploaded to {saved_path}: {att.filename}"
                        "  -- pip install pdfplumber to extract text, or use filesystem tools]"
                    )
                    continue
            elif ext == ".docx":
                extracted = extract_docx_text(att.data)
                if extracted is None:
                    ref_notes.append(
                        f"[DOCX uploaded to {saved_path}: {att.filename}"
                        "  -- pip install python-docx to extract text, or use filesystem tools]"
                    )
                    continue

            if extracted:
                block_text = f"**{att.filename}** (extracted text)\n```\n{extracted}\n```"
                blocks.append({"type": "text", "text": block_text})
                inline_count += 1
                inline_bytes += len(att.data)
            else:
                ref_notes.append(
                    f"[Document uploaded to {saved_path}: {att.filename}]"
                )

    # --- assemble final content ---
    if not blocks and not ref_notes:
        return text

    full_text = text
    if ref_notes:
        full_text = (text + "\n\n" + "\n".join(ref_notes)).strip()

    if not blocks:
        # All reference  --" return plain string, no content block list needed
        return full_text

    # Mix: text block first, then attachment blocks
    result: list[dict] = [{"type": "text", "text": full_text}] + blocks
    return result
